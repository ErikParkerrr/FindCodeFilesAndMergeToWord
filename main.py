import os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

FILE_EXTENSION = ".vhd"

def remove_header_comments(vhdl_code):
   # Find the start and end markers
    start_marker = "----------------------------------------------------------------------------------"
    end_marker = "----------------------------------------------------------------------------------"

    # Find the indices of the start and end markers
    start_index = vhdl_code.find(start_marker)
    end_index = vhdl_code.find(end_marker, start_index + len(start_marker))

    # Remove the text between the markers
    if start_index != -1 and end_index != -1:
        vhdl_code = vhdl_code[:start_index] + vhdl_code[end_index + len(end_marker):]

    return vhdl_code

def add_vhdl_to_docx(document, vhdl_file):
    # Add file name as a heading
    document.add_heading(f'{os.path.basename(vhdl_file)}', level=2).paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Read VHDL code
    with open(vhdl_file, 'r') as file:
        vhdl_code = file.read()

    # Remove header comments block
    vhdl_code = remove_header_comments(vhdl_code)

    # Create a new paragraph for VHDL code
    paragraph = document.add_paragraph(vhdl_code, style='Code')

    # Set font size
    font = paragraph.runs[0].font
    font.size = Pt(10)

def merge_vhdl_files(output_file_name="merged_vhdl.docx"):
    # Get the current directory
    current_directory = os.path.dirname(os.path.abspath(__file__))

    # Initialize a list to store the paths of all VHDL files
    vhdl_files = []

    # Recursively search for VHDL files in all subdirectories
    for root, dirs, files in os.walk(current_directory):
        vhdl_files.extend([os.path.join(root, file) for file in files if file.endswith(FILE_EXTENSION)])

    # Check if there are any VHDL files
    if not vhdl_files:
        print("No VHDL files found in the current directory or its subdirectories.")
        return

    # Create a new Word document
    document = Document()

    # Add a custom style for VHDL code
    style = document.styles.add_style('Code', WD_PARAGRAPH_ALIGNMENT.CENTER)
    font = style.font
    font.name = 'Courier New'
    font.size = Pt(10)

    # Loop through each VHDL file and add its content to the document
    for vhdl_file in vhdl_files:
        add_vhdl_to_docx(document, vhdl_file)

    # Save the document
    document.save(output_file_name)

    print(f"Merge complete. Merged content saved in {output_file_name}.")

# Call the function to merge VHDL files
merge_vhdl_files()
